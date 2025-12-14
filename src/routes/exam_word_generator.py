"""
نظام توليد ملفات Word لنماذج الاختبار
يقوم بإنشاء ملفات Word قابلة للتعديل مع تنسيق احترافي
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io


class ExamWordGenerator:
    """فئة لتوليد ملفات Word لنماذج الاختبار"""
    
    def __init__(self, exam_title="نموذج الاختبار", include_answers=False):
        """
        تهيئة منشئ ملفات Word
        
        Args:
            exam_title: عنوان الاختبار
            include_answers: هل تتضمن الملف الإجابات
        """
        self.document = Document()
        self.exam_title = exam_title
        self.include_answers = include_answers
        self.setup_document()
    
    def setup_document(self):
        """إعداد خصائص الوثيقة الأساسية"""
        # تعيين اتجاه النص من اليمين إلى اليسار
        sections = self.document.sections
        for section in sections:
            section.right_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
    
    def add_header(self, course_name="", unit_name="", lesson_name="", date=None):
        """
        إضافة رأس الصفحة (الكليشه)
        
        Args:
            course_name: اسم المنهج
            unit_name: اسم الوحدة
            lesson_name: اسم الدرس
            date: التاريخ (إذا لم يتم تحديده، سيتم استخدام التاريخ الحالي)
        """
        if date is None:
            date = datetime.now().strftime("%d/%m/%Y")
        
        # العنوان الرئيسي
        title = self.document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(self.exam_title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # إضافة خط فاصل
        self._add_horizontal_line()
        
        # معلومات الكليشه
        header_table = self.document.add_table(rows=4, cols=2)
        header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # تعيين عرض الأعمدة
        for row in header_table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(2)
        
        # ملء بيانات الكليشه
        header_data = [
            ("المنهج:", course_name or "_______________"),
            ("الوحدة:", unit_name or "_______________"),
            ("الدرس:", lesson_name or "_______________"),
            ("التاريخ:", date)
        ]
        
        for i, (label, value) in enumerate(header_data):
            row = header_table.rows[i]
            
            # الخلية اليمنى (التسمية)
            label_cell = row.cells[1]
            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label_run = label_para.add_run(label)
            label_run.font.bold = True
            label_run.font.size = Pt(11)
            
            # الخلية اليسرى (القيمة)
            value_cell = row.cells[0]
            value_para = value_cell.paragraphs[0]
            value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            value_run = value_para.add_run(value)
            value_run.font.size = Pt(11)
        
        # إضافة مسافة
        self.document.add_paragraph()
    
    def add_questions(self, questions):
        """
        إضافة الأسئلة إلى الوثيقة
        
        Args:
            questions: قائمة الأسئلة (قاموس يحتوي على question_text, options, correct_option_id)
        """
        for idx, question in enumerate(questions, 1):
            # رقم السؤال
            q_number = self.document.add_paragraph()
            q_number.paragraph_format.space_before = Pt(12)
            q_number.paragraph_format.space_after = Pt(6)
            
            q_number_run = q_number.add_run(f"السؤال {idx}: ")
            q_number_run.font.bold = True
            q_number_run.font.size = Pt(12)
            q_number_run.font.color.rgb = RGBColor(102, 126, 234)
            
            # نص السؤال
            q_text_run = q_number.add_run(question.get('question_text', ''))
            q_text_run.font.size = Pt(11)
            
            # الخيارات
            options = question.get('options', [])
            correct_option_id = question.get('correct_option_id')
            
            for opt_idx, option in enumerate(options):
                opt_para = self.document.add_paragraph()
                opt_para.paragraph_format.right_indent = Inches(0.5)
                opt_para.paragraph_format.space_before = Pt(3)
                opt_para.paragraph_format.space_after = Pt(3)
                
                # الحرف (أ، ب، ج، د)
                opt_letter = chr(ord('أ') + opt_idx)
                opt_letter_run = opt_para.add_run(f"{opt_letter}) ")
                opt_letter_run.font.bold = True
                opt_letter_run.font.size = Pt(11)
                
                # نص الخيار
                opt_text_run = opt_para.add_run(option.get('option_text', ''))
                opt_text_run.font.size = Pt(11)
                
                # إذا كان هذا الخيار صحيحاً وكنا نتضمن الإجابات
                if self.include_answers and option.get('option_id') == correct_option_id:
                    opt_text_run.font.bold = True
                    opt_text_run.font.color.rgb = RGBColor(40, 167, 69)  # أخضر
                    
                    # إضافة علامة الصحة
                    check_run = opt_para.add_run(" ✓")
                    check_run.font.bold = True
                    check_run.font.color.rgb = RGBColor(40, 167, 69)
            
            # إضافة خط فاصل بين الأسئلة
            if idx < len(questions):
                self._add_horizontal_line(style='dashed')
    
    def add_answer_key(self, questions):
        """
        إضافة نموذج الإجابة
        
        Args:
            questions: قائمة الأسئلة
        """
        if not self.include_answers:
            return
        
        # إضافة فاصل صفحة
        self.document.add_page_break()
        
        # عنوان نموذج الإجابة
        answer_title = self.document.add_paragraph()
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        answer_title.paragraph_format.space_before = Pt(12)
        answer_title.paragraph_format.space_after = Pt(12)
        
        title_run = answer_title.add_run("نموذج الإجابة")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # إضافة خط فاصل
        self._add_horizontal_line()
        
        # جدول الإجابات
        answers_table = self.document.add_table(rows=1, cols=5)
        answers_table.style = 'Light Grid Accent 1'
        
        # رأس الجدول
        header_cells = answers_table.rows[0].cells
        headers = ["السؤال 5-1", "السؤال 10-6", "السؤال 15-11", "السؤال 20-16", "السؤال 25-21"]
        
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            # تنسيق رأس الجدول
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # تلوين رأس الجدول
            self._shade_cell(cell, "4472C4")
        
        # ملء الإجابات
        answers_row = answers_table.add_row()
        
        for col_idx in range(5):
            cell = answers_row.cells[col_idx]
            cell.text = ""
            
            for q_idx in range(col_idx * 5, min((col_idx + 1) * 5, len(questions))):
                if q_idx < len(questions):
                    question = questions[q_idx]
                    correct_option_id = question.get('correct_option_id')
                    
                    # البحث عن الحرف الصحيح
                    for opt_idx, option in enumerate(question.get('options', [])):
                        if option.get('option_id') == correct_option_id:
                            answer_letter = chr(ord('أ') + opt_idx)
                            break
                    else:
                        answer_letter = "-"
                    
                    para = cell.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run(f"{q_idx + 1}: {answer_letter}")
                    run.font.size = Pt(11)
                    run.font.bold = True
    
    def _add_horizontal_line(self, style='single'):
        """
        إضافة خط فاصل أفقي
        
        Args:
            style: نمط الخط ('single' أو 'dashed')
        """
        para = self.document.add_paragraph()
        pPr = para._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), style)
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _shade_cell(self, cell, color):
        """
        تلوين خلية الجدول
        
        Args:
            cell: الخلية
            color: اللون (بصيغة hex بدون #)
        """
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def save(self, filename):
        """
        حفظ الوثيقة في ملف
        
        Args:
            filename: اسم الملف
        """
        self.document.save(filename)
    
    def get_bytes(self):
        """
        الحصول على محتوى الوثيقة كـ bytes
        
        Returns:
            محتوى الوثيقة كـ bytes
        """
        bytes_io = io.BytesIO()
        self.document.save(bytes_io)
        bytes_io.seek(0)
        return bytes_io.getvalue()


def generate_exam_word(questions, exam_title="نموذج الاختبار", 
                       course_name="", unit_name="", lesson_name="",
                       include_answers=False, filename=None):
    """
    دالة مساعدة لتوليد ملف Word للاختبار
    
    Args:
        questions: قائمة الأسئلة
        exam_title: عنوان الاختبار
        course_name: اسم المنهج
        unit_name: اسم الوحدة
        lesson_name: اسم الدرس
        include_answers: هل تتضمن الملف الإجابات
        filename: اسم الملف (إذا كان None، سيتم إرجاع bytes)
    
    Returns:
        اسم الملف أو bytes حسب المعامل filename
    """
    generator = ExamWordGenerator(exam_title, include_answers)
    generator.add_header(course_name, unit_name, lesson_name)
    generator.add_questions(questions)
    generator.add_answer_key(questions)
    
    if filename:
        generator.save(filename)
        return filename
    else:
        return generator.get_bytes()


# مثال على الاستخدام
if __name__ == "__main__":
    # بيانات تجريبية
    sample_questions = [
        {
            'question_id': 1,
            'question_text': 'ما هو العدد الذري للكربون؟',
            'options': [
                {'option_id': 1, 'option_text': '4'},
                {'option_id': 2, 'option_text': '6'},
                {'option_id': 3, 'option_text': '8'},
                {'option_id': 4, 'option_text': '12'}
            ],
            'correct_option_id': 2
        },
        {
            'question_id': 2,
            'question_text': 'ما هو الرمز الكيميائي للذهب؟',
            'options': [
                {'option_id': 1, 'option_text': 'Go'},
                {'option_id': 2, 'option_text': 'Gd'},
                {'option_id': 3, 'option_text': 'Au'},
                {'option_id': 4, 'option_text': 'Ag'}
            ],
            'correct_option_id': 3
        }
    ]
    
    # توليد ملف Word بدون إجابات
    generate_exam_word(
        sample_questions,
        exam_title="نموذج الاختبار - الكيمياء",
        course_name="الكيمياء العامة",
        unit_name="الوحدة الأولى",
        lesson_name="الذرات والجزيئات",
        include_answers=False,
        filename="exam_without_answers.docx"
    )
    
    # توليد ملف Word مع الإجابات
    generate_exam_word(
        sample_questions,
        exam_title="نموذج الاختبار - الكيمياء",
        course_name="الكيمياء العامة",
        unit_name="الوحدة الأولى",
        lesson_name="الذرات والجزيئات",
        include_answers=True,
        filename="exam_with_answers.docx"
    )
    
    print("✓ تم توليد ملفات Word بنجاح")
