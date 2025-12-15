"""
نظام محسّن لتوليد ملفات PDF و Word من بيانات الاختبار
مع شعار وزارة التعليم والتصميم الاحترافي
"""

from jinja2 import Template
from datetime import datetime
import io
import base64
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image


class ExamGenerator:
    """فئة محسّنة لتوليد الاختبارات بصيغ مختلفة"""
    
    def __init__(self, header_settings=None, logo_path=None):
        """تهيئة منشئ الاختبارات"""
        self.header_settings = header_settings or {}
        self.logo_path = logo_path or '/home/ubuntu/ministry_logo.png'
        self.html_template = self._get_html_template()
    
    def _get_logo_base64(self):
        """تحويل الشعار إلى base64"""
        try:
            with open(self.logo_path, 'rb') as f:
                return base64.b64encode(f.read()).decode()
        except:
            return None
    
    def _get_html_template(self):
        """الحصول على قالب HTML"""
        return """
<!DOCTYPE html>
<html dir="rtl" lang="ar">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ exam_title }}</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Arial', sans-serif;
            direction: rtl;
            line-height: 1.6;
            background-color: #fff;
            color: #333;
        }
        
        .container {
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
        }
        
        /* رأس الاختبار */
        .header {
            border: 2px solid #000;
            padding: 20px;
            margin-bottom: 20px;
            text-align: center;
        }
        
        .logo-container {
            text-align: center;
            margin-bottom: 15px;
        }
        
        .logo-container img {
            height: 70px;
            width: auto;
        }
        
        .header-info {
            display: grid;
            grid-template-columns: 1fr 1fr 1fr;
            gap: 10px;
            margin-bottom: 15px;
            text-align: right;
            font-size: 13px;
            line-height: 1.8;
        }
        
        .header-info div {
            text-align: right;
        }
        
        .header-table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 10px;
        }
        
        .header-table td {
            border: 1px solid #000;
            padding: 8px;
            text-align: right;
            font-size: 12px;
        }
        
        .header-table .label {
            font-weight: bold;
            background-color: #f5f5f5;
            width: 25%;
        }
        
        .header-table .value {
            width: 25%;
        }
        
        .exam-title {
            font-size: 16px;
            font-weight: bold;
            margin: 15px 0;
            border-bottom: 2px solid #000;
            padding-bottom: 10px;
        }
        
        /* الأسئلة */
        .questions {
            margin-top: 20px;
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
        }
        
        .question {
            margin-bottom: 20px;
            page-break-inside: avoid;
            border: 1px solid #ccc;
            padding: 15px;
            border-radius: 5px;
            background-color: #f9f9f9;
        }
        
        .question-number {
            font-weight: bold;
            margin-bottom: 5px;
            color: #0066cc;
            font-size: 12px;
        }
        
        .question-text {
            margin-bottom: 10px;
            line-height: 1.8;
            font-weight: 500;
            font-size: 12px;
        }
        
        .options {
            margin-right: 0;
            margin-top: 10px;
        }
        
        .options-row {
            display: flex;
            justify-content: space-between;
            gap: 60px;
            margin-bottom: 8px;
            flex-wrap: nowrap;
            text-align: right;
        }
        
        .option {
            margin-bottom: 0;
            line-height: 1.6;
            font-size: 11px;
            display: flex;
            align-items: flex-start;
            gap: 5px;
            flex: 0 1 auto;
        }
        
        .option span {
            flex-shrink: 0;
            min-width: 20px;
        }
        
        /* جدول الإجابات */
        .answer-key {
            margin-top: 30px;
            border-top: 2px solid #000;
            padding-top: 20px;
            grid-column: 1 / -1;
        }
        
        .answer-table {
            width: 100%;
            border-collapse: collapse;
        }
        
        .answer-table th,
        .answer-table td {
            border: 1px solid #000;
            padding: 8px;
            text-align: center;
            font-size: 11px;
        }
        
        .answer-table th {
            background-color: #f5f5f5;
            font-weight: bold;
        }
        
        /* جدول المعلومات الإضافية */
        .info-table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 10px;
        }
        
        .info-table td {
            border: 1px solid #000;
            padding: 8px;
            text-align: right;
            font-size: 12px;
        }
        
        .info-table .label {
            font-weight: bold;
            background-color: #f5f5f5;
            width: 25%;
        }
        
        .info-table .value {
            width: 25%;
        }
        
        .student-info {
            margin-top: 15px;
            display: grid;
            grid-template-columns: 1fr 1fr 1fr;
            gap: 20px;
            font-size: 12px;
            line-height: 2;
        }
        
        .student-info-item {
            text-align: right;
            border-bottom: 1px solid #000;
        }
        
        .header-top {
            display: grid;
            grid-template-columns: 1fr 1fr 1fr;
            gap: 20px;
            margin-bottom: 20px;
            font-size: 14px;
            line-height: 1.8;
        }
        
        .header-left {
            text-align: right;
        }
        
        .header-center {
            text-align: center;
        }
        
        .header-right {
            text-align: right;
        }
        
        .info-section {
            margin: 10px 0;
            padding: 10px 0;
            text-align: right;
            font-size: 12px;
            line-height: 1.8;
        }
    </style>
</head>
<body>
    <div class="container">
        <!-- رأس الاختبار -->
        <div class="header">
            <div class="header-top">
                <div class="header-left">
                    <div><strong>المادة:</strong> {{ subject }}</div>
                    <div><strong>الزمن:</strong> {{ time }}</div>
                    <div><strong>الصف:</strong> {{ grade }}</div>
                </div>
                <div class="header-center">
                    <div><strong>{{ country }}</strong></div>
                    <div><strong>{{ ministry }}</strong></div>
                </div>
                <div class="header-right">
                    <div><strong>{{ education_department }}</strong></div>
                    <div><strong>{{ school_name }}</strong></div>
                </div>
            </div>
            
            <div class="exam-title">{{ exam_title }}</div>
            
            <table class="header-table">
                <tr>
                    <td class="label">المراجع</td>
                    <td class="value"></td>
                    <td class="label">المصحح</td>
                    <td class="value"></td>
                    <td class="label">درجة كتابية</td>
                    <td class="value"></td>
                    <td class="label">درجة الطالب رقماً</td>
                    <td class="value"></td>
                    <td class="label">الدرجة الأساسية</td>
                    <td class="value">{{ total_score }}</td>
                </tr>
                <tr>
                    <td class="label"></td>
                    <td class="value"></td>
                    <td class="label"></td>
                    <td class="value"></td>
                    <td class="label"></td>
                    <td class="value"></td>
                    <td class="label"></td>
                    <td class="value"></td>
                    <td class="label"></td>
                    <td class="value"></td>
                </tr>
            </table>
            
            <div class="student-info">
                <div class="student-info-item">اسم الطالب: __________</div>
                <div class="student-info-item">الشعبة: __________</div>
                <div class="student-info-item">رقم الجلوس: __________</div>
            </div>
        </div>
        
        <!-- الأسئلة -->
        <div class="questions">
            {% for question in questions %}
            <div class="question">
                <div class="question-number">السؤال {{ loop.index }}: ({{ question.points }} درجات)</div>
                <div class="question-text">{{ question.question_text }}</div>
                <div class="options">
                    {% set letters = ['أ', 'ب', 'ج', 'د'] %}
                    {% for row_idx in range(0, question.options|length, 2) %}
                    <div class="options-row">
                        {% for col_idx in range(2) %}
                            {% if row_idx + col_idx < question.options|length %}
                            <div class="option">
                                <span><strong>{{ letters[row_idx + col_idx] if row_idx + col_idx < 4 else row_idx + col_idx + 1 }})&nbsp;</strong></span>
                                {{ question.options[row_idx + col_idx].option_text }}
                            </div>
                            {% endif %}
                        {% endfor %}
                    </div>
                    {% endfor %}
                </div>
            </div>
            {% endfor %}
            
            {% if show_answers %}
            <div class="answer-key">
                <div style="font-weight: bold; margin-bottom: 10px;">مفتاح الإجابات</div>
                <table class="answer-table">
                    <tr>
                        {% for i in range(1, questions|length + 1) %}
                        <th>{{ i }}</th>
                        {% endfor %}
                    </tr>
                    <tr>
                        {% for question in questions %}
                        <td>{{ question.correct_answer }}</td>
                        {% endfor %}
                    </tr>
                </table>
            </div>
            {% endif %}
        </div>
    </div>
</body>
</html>
"""
    
    def _prepare_context(self, questions, exam_title, show_answers, **kwargs):
        """تحضير السياق للقالب"""
        context = {
            'exam_title': exam_title,
            'country': kwargs.get('country', self.header_settings.get('country', 'المملكة العربية السعودية')),
            'ministry': kwargs.get('ministry', self.header_settings.get('ministry', 'وزارة التعليم')),
            'education_department': kwargs.get('education_department', self.header_settings.get('education_department', 'الإدارة العامة للتعليم بالمنطقة الشرقية')),
            'school_name': kwargs.get('school_name', self.header_settings.get('school_name', 'مدرسة عبدالرحمن بن القاسم الثانوية')),
            'subject': kwargs.get('subject', self.header_settings.get('subject', 'كيمياء 4')),
            'time': kwargs.get('time', self.header_settings.get('time', 'ثلاث ساعات')),
            'grade': kwargs.get('grade', self.header_settings.get('grade', 'ثالث ثانوي')),
            'total_score': kwargs.get('total_score', self.header_settings.get('total_score', 30)),
            'questions': [],
            'show_answers': show_answers,
            'logo': self._get_logo_base64()
        }
        
        # تنسيق الأسئلة
        letters = ['أ', 'ب', 'ج', 'د']
        
        for question in questions:
            formatted_q = {
                'question_text': question.get('question_text', ''),
                'points': question.get('points', 1),
                'options': [],
                'correct_answer': ''
            }
            
            options = question.get('options', [])
            for idx, option in enumerate(options):
                formatted_q['options'].append({
                    'letter': letters[idx] if idx < len(letters) else str(idx + 1),
                    'option_text': option.get('option_text', '')
                })
                
                if option.get('is_correct') or option.get('option_id') == question.get('correct_option_id'):
                    formatted_q['correct_answer'] = letters[idx] if idx < len(letters) else str(idx + 1)
            
            context['questions'].append(formatted_q)
        
        return context
    
    def generate_html(self, questions, exam_title="نموذج الاختبار", 
                     show_answers=False, **kwargs):
        """توليد HTML من البيانات"""
        context = self._prepare_context(questions, exam_title, show_answers, **kwargs)
        template = Template(self.html_template)
        return template.render(**context)
    
    def generate_pdf(self, questions, exam_title="نموذج الاختبار", 
                    show_answers=False, **kwargs):
        """توليد PDF من HTML"""
        html_content = self.generate_html(questions, exam_title, show_answers, **kwargs)
        html_obj = HTML(string=html_content)
        pdf_bytes = html_obj.write_pdf()
        return pdf_bytes
    
    def generate_word(self, questions, exam_title="نموذج الاختبار", 
                     show_answers=False, **kwargs):
        """توليد Word بنفس تصميم PDF"""
        try:
            # إنشاء مستند Word
            doc = Document()
            
            # تعيين هوامش الصفحة
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # إضافة الشعار
            try:
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                run.add_picture(self.logo_path, width=Inches(1.2))
            except:
                pass
            
            # إضافة معلومات الرأس
            header_data = {
                'country': kwargs.get('country', self.header_settings.get('country', 'المملكة العربية السعودية')),
                'ministry': kwargs.get('ministry', self.header_settings.get('ministry', 'وزارة التعليم')),
                'education_department': kwargs.get('education_department', self.header_settings.get('education_department', 'الإدارة العامة للتعليم بالمنطقة الشرقية')),
                'school_name': kwargs.get('school_name', self.header_settings.get('school_name', 'مدرسة عبدالرحمن بن القاسم الثانوية')),
                'subject': kwargs.get('subject', self.header_settings.get('subject', 'كيمياء 4')),
                'time': kwargs.get('time', self.header_settings.get('time', 'ثلاث ساعات')),
                'grade': kwargs.get('grade', self.header_settings.get('grade', 'ثالث ثانوي')),
                'total_score': kwargs.get('total_score', self.header_settings.get('total_score', 30)),
            }
            
            # جدول معلومات الرأس
            header_info = doc.add_table(rows=1, cols=3)
            header_info.style = 'Table Grid'
            
            cells = header_info.rows[0].cells
            
            # العمود الأيمن
            right_cell = cells[2]
            right_cell.text = f"{header_data['country']}\n{header_data['ministry']}"
            for paragraph in right_cell.paragraphs:
                paragraph.paragraph_format.direction = 1
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
            
            # العمود الأوسط (فارغ)
            middle_cell = cells[1]
            middle_cell.text = ""
            
            # العمود الأيسر
            left_cell = cells[0]
            left_cell.text = f"{header_data['education_department']}\n{header_data['school_name']}"
            for paragraph in left_cell.paragraphs:
                paragraph.paragraph_format.direction = 1
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
            
            # عنوان الاختبار
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.direction = 1
            title.paragraph_format.space_before = Pt(12)
            title.paragraph_format.space_after = Pt(12)
            title_run = title.add_run(exam_title)
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            
            # إضافة خط تحت العنوان
            pPr = title._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '24')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
            
            # جدول المعلومات
            info_table = doc.add_table(rows=2, cols=4)
            info_table.style = 'Table Grid'
            
            info_cells = [
                ('المادة:', header_data['subject'], 'الزمن:', header_data['time']),
                ('الصف:', header_data['grade'], 'الدرجة الكلية:', str(header_data['total_score'])),
            ]
            
            for row_idx, row_data in enumerate(info_cells):
                row = info_table.rows[row_idx]
                for col_idx, (label, value) in enumerate([(row_data[0], row_data[1]), (row_data[2], row_data[3])]):
                    label_cell = row.cells[col_idx * 2]
                    label_cell.text = label
                    for paragraph in label_cell.paragraphs:
                        paragraph.paragraph_format.direction = 1
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                            run.font.bold = True
                        self._add_shading(paragraph, 'f5f5f5')
                    
                    value_cell = row.cells[col_idx * 2 + 1]
                    value_cell.text = value
                    for paragraph in value_cell.paragraphs:
                        paragraph.paragraph_format.direction = 1
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
            
            # جدول الدرجات
            grades_table = doc.add_table(rows=2, cols=5)
            grades_table.style = 'Table Grid'
            
            grade_headers = ['الدرجة الأساسية', 'درجة الطالب رقماً', 'درجة كتابية', 'المصحح', 'المراجع']
            for col_idx, header in enumerate(grade_headers):
                cell = grades_table.rows[0].cells[col_idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.direction = 1
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                    self._add_shading(paragraph, 'f5f5f5')
            
            # ملء الصف الثاني من جدول الدرجات
            grades_table.rows[1].cells[0].text = str(header_data['total_score'])
            
            # معلومات الطالب
            doc.add_paragraph()
            student_info = doc.add_table(rows=1, cols=3)
            student_info.style = 'Table Grid'
            
            student_fields = [
                'اسم الطالب: __________',
                'الشعبة: __________',
                'رقم الجلوس: __________'
            ]
            
            for col_idx, field in enumerate(student_fields):
                cell = student_info.rows[0].cells[col_idx]
                cell.text = field
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.direction = 1
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
            
            # إضافة فاصل
            doc.add_paragraph()
            
            # إضافة الأسئلة في تخطيط عمودين
            letters = ['أ', 'ب', 'ج', 'د']
            num_rows = (len(questions) + 1) // 2
            
            questions_table = doc.add_table(rows=num_rows, cols=2)
            questions_table.style = 'Table Grid'
            
            for row in questions_table.rows:
                row.cells[0].width = Inches(3.5)
                row.cells[1].width = Inches(3.5)
            
            for row_idx in range(num_rows):
                row = questions_table.rows[row_idx]
                
                right_cell = row.cells[1]
                right_q_idx = row_idx
                
                left_cell = row.cells[0]
                left_q_idx = row_idx + num_rows
                
                if right_q_idx < len(questions):
                    self._fill_question_cell(right_cell, questions[right_q_idx], right_q_idx + 1, letters)
                
                if left_q_idx < len(questions):
                    self._fill_question_cell(left_cell, questions[left_q_idx], left_q_idx + 1, letters)
            
            # جدول الإجابات
            if show_answers:
                doc.add_paragraph()
                answer_key_para = doc.add_paragraph("مفتاح الإجابات")
                answer_key_para.paragraph_format.direction = 1
                answer_key_para.runs[0].bold = True
                answer_key_para.runs[0].font.size = Pt(12)
                
                answer_table = doc.add_table(rows=2, cols=len(questions))
                answer_table.style = 'Table Grid'
                
                for i in range(len(questions)):
                    cell = answer_table.rows[0].cells[i]
                    cell.text = str(i + 1)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.direction = 1
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                        self._add_shading(paragraph, 'f5f5f5')
                
                for i, question in enumerate(questions):
                    options = question.get('options', [])
                    correct_answer = ''
                    for opt_idx, option in enumerate(options):
                        if option.get('is_correct') or option.get('option_id') == question.get('correct_option_id'):
                            correct_answer = letters[opt_idx] if opt_idx < len(letters) else str(opt_idx + 1)
                            break
                    cell = answer_table.rows[1].cells[i]
                    cell.text = correct_answer
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.direction = 1
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # حفظ في BytesIO
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            return doc_bytes.getvalue()
        
        except Exception as e:
            raise Exception(f"خطأ في توليد ملف Word: {str(e)}")
    
    def _add_shading(self, paragraph, color):
        """إضافة خلفية للفقرة"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        paragraph._element.get_or_add_pPr().append(shading_elm)
    
    def _fill_question_cell(self, cell, question, question_num, letters):
        """ملء خلية السؤال"""
        for paragraph in cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        q_num = cell.add_paragraph()
        q_num.paragraph_format.direction = 1
        q_num.paragraph_format.space_after = Pt(6)
        q_num_run = q_num.add_run(f"السؤال {question_num}: ({question.get('points', 1)} درجات)")
        q_num_run.font.bold = True
        q_num_run.font.size = Pt(11)
        q_num_run.font.color.rgb = RGBColor(0, 102, 204)
        
        q_text = cell.add_paragraph(question.get('question_text', ''))
        q_text.paragraph_format.direction = 1
        q_text.paragraph_format.right_indent = Inches(0.1)
        q_text.paragraph_format.space_after = Pt(6)
        for run in q_text.runs:
            run.font.size = Pt(10)
        
        options = question.get('options', [])
        if len(options) > 0:
            for row_idx in range(0, len(options), 2):
                opt_para = cell.add_paragraph()
                opt_para.paragraph_format.direction = 1
                opt_para.paragraph_format.space_after = Pt(4)
                
                if row_idx < len(options):
                    letter = letters[row_idx] if row_idx < len(letters) else str(row_idx + 1)
                    letter_run = opt_para.add_run(f"{letter}) ")
                    letter_run.font.bold = True
                    letter_run.font.size = Pt(10)
                    opt_run = opt_para.add_run(options[row_idx].get('option_text', ''))
                    opt_run.font.size = Pt(10)
                    
                    space_run = opt_para.add_run("                    ")
                    space_run.font.size = Pt(10)
                
                if row_idx + 1 < len(options):
                    letter = letters[row_idx + 1] if row_idx + 1 < len(letters) else str(row_idx + 2)
                    letter_run = opt_para.add_run(f"{letter}) ")
                    letter_run.font.bold = True
                    letter_run.font.size = Pt(10)
                    opt_run = opt_para.add_run(options[row_idx + 1].get('option_text', ''))
                    opt_run.font.size = Pt(10)


# دالة موحدة
def generate_exam(questions, exam_title="نموذج الاختبار", 
                 output_format='word', show_answers=False, 
                 header_settings=None, logo_path=None, **kwargs):
    """دالة موحدة لتوليد الاختبارات"""
    generator = ExamGenerator(header_settings, logo_path)
    
    if output_format == 'pdf':
        return generator.generate_pdf(questions, exam_title, show_answers, **kwargs)
    else:
        return generator.generate_word(questions, exam_title, show_answers, **kwargs)
